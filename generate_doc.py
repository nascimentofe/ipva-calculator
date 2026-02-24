from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(30, 41, 59)

# --- Title ---
title = doc.add_heading('Calculadora do Novo IPVA', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PEC do Teto de 1% — Documentação Técnica do Projeto')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()

# --- 1. Visão Geral ---
doc.add_heading('1. Visão Geral do Projeto', level=1)
doc.add_paragraph(
    'Sistema web de página única (SPA estática) que permite ao usuário calcular '
    'a economia no IPVA caso a PEC que propõe teto constitucional de 1% seja aprovada. '
    'O sistema compara a alíquota atual do estado do usuário com o novo teto proposto, '
    'projetando economia anual e em 5 anos.'
)
doc.add_paragraph(
    'O projeto é composto por HTML, CSS e JavaScript puros (sem frameworks), '
    'garantindo performance máxima, SEO nativo e hospedagem como site estático no Render.'
)

# --- 2. Estrutura ---
doc.add_heading('2. Estrutura de Arquivos', level=1)
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Arquivo'
hdr[1].text = 'Descrição'
data = [
    ('index.html', 'Página principal com formulário, resultados, ads AdSense e meta tags SEO'),
    ('style.css', 'Estilos responsivos com design moderno, ads e cookie consent'),
    ('script.js', 'Lógica de cálculo, máscara monetária, estados, analytics e consentimento LGPD'),
    ('ads.txt', 'Arquivo de verificação obrigatório do Google AdSense'),
    ('render.yaml', 'Configuração de deploy como Static Site no Render'),
]
for i, (arq, desc) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = arq
    row[1].text = desc

doc.add_paragraph()

# --- 3. Alíquotas ---
doc.add_heading('3. Alíquotas de IPVA por Estado (Veículos de Passeio)', level=1)
doc.add_paragraph(
    'As alíquotas abaixo são referentes a automóveis de passeio particulares, '
    'conforme dados oficiais das Secretarias da Fazenda estaduais (2025/2026).'
)

aliquotas = [
    ('4,0%', 'SP, MG, RJ'),
    ('3,75%', 'GO'),
    ('3,5%', 'DF, PR'),
    ('3,0%', 'AL, AM, AP, MS, MT, PE, RN, RO, RR, RS'),
    ('2,5%', 'BA, CE, MA, PA, PB, PI, SE'),
    ('2,0%', 'AC, ES, SC, TO'),
]
t2 = doc.add_table(rows=len(aliquotas) + 1, cols=2)
t2.style = 'Light Grid Accent 1'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.rows[0].cells[0].text = 'Alíquota'
t2.rows[0].cells[1].text = 'Estados'
for i, (aliq, estados) in enumerate(aliquotas, 1):
    t2.rows[i].cells[0].text = aliq
    t2.rows[i].cells[1].text = estados

doc.add_paragraph()

# --- 4. Lógica de Cálculo ---
doc.add_heading('4. Lógica de Cálculo', level=1)

doc.add_heading('4.1. Fórmula do IPVA Atual', level=2)
p = doc.add_paragraph()
run = p.add_run('IPVA Atual = Valor FIPE × (Alíquota do Estado / 100)')
run.bold = True
doc.add_paragraph(
    'Exemplo: Veículo de R$ 100.000 em SP (4%) → IPVA = R$ 4.000,00'
)

doc.add_heading('4.2. Fórmula do Novo IPVA (PEC)', level=2)
p = doc.add_paragraph()
run = p.add_run('IPVA Novo = Valor FIPE × (Teto de 1% / 100)')
run.bold = True
doc.add_paragraph(
    'Exemplo: Veículo de R$ 100.000 com teto de 1% → IPVA = R$ 1.000,00'
)

doc.add_heading('4.3. Cálculo por Peso (Opcional)', level=2)
doc.add_paragraph(
    'Se o peso do veículo for informado, aplica-se um fator proporcional ao teto de 1%:'
)
p = doc.add_paragraph()
run = p.add_run('Fator Peso = max(0.3, min(1.0, peso_kg / 2000))')
run.bold = True
doc.add_paragraph(
    'Isso significa que veículos com até 600kg pagam 0,3% e veículos com 2000kg+ pagam 1,0%. '
    'Veículos intermediários têm alíquota proporcional.'
)

doc.add_heading('4.4. Projeções', level=2)
bullets = [
    'Economia Anual = IPVA Atual − IPVA Novo',
    'Redução (%) = (Economia Anual / IPVA Atual) × 100',
    'Economia em 5 Anos = Economia Anual × 5',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()

# --- 5. Funcionalidades ---
doc.add_heading('5. Funcionalidades do Sistema', level=1)
features = [
    ('Formulário Principal', 'Seleção de estado (com alíquota exibida), valor FIPE com máscara monetária, peso opcional'),
    ('Resultado Visual', 'Cards comparativos (atual vs novo), economia anual, redução percentual, economia em 5 anos'),
    ('Compartilhamento', 'Botão de compartilhar via Web Share API ou clipboard'),
    ('FAQ', 'Perguntas frequentes sobre a PEC com detalhes expandíveis'),
    ('Responsividade', 'Layout adaptável para desktop, tablet e mobile'),
    ('SEO', 'Meta tags Open Graph, Twitter Cards, canonical URL, schema semântico'),
]
for feat, desc in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(feat + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

# --- 6. Ads (Google AdSense) ---
doc.add_heading('6. Publicidade — Google AdSense', level=1)

doc.add_heading('6.1. Pré-requisitos para Aprovação do AdSense', level=2)
prereqs = [
    'Domínio próprio ativo (ex: ipva.fsncompany.com.br) — AdSense não aceita subdomínios do Render (.onrender.com)',
    'Conteúdo original e útil publicado (a calculadora já atende esse requisito)',
    'Política de Privacidade publicada no site (obrigatório pela LGPD e pelo AdSense)',
    'Mínimo de 2-4 semanas de site no ar com tráfego antes de solicitar aprovação',
    'O site não pode ter conteúdo adulto, violento ou ilegal',
]
for item in prereqs:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2. Passo a Passo para Integração', level=2)
steps_ads = [
    'Acesse adsense.google.com e clique em "Começar"',
    'Cadastre o domínio "ipva.fsncompany.com.br" e aceite os termos',
    'O Google fornecerá um Publisher ID no formato "ca-pub-XXXXXXXXXXXXXXXX"',
    'No index.html, substitua TODAS as ocorrências de "ca-pub-XXXXXXXXXXXXXXXX" pelo seu Publisher ID real',
    'O Google fornecerá também Slot IDs para cada unidade de anúncio — substitua os valores _AD_SLOT_ID correspondentes',
    'No arquivo ads.txt (na raiz do site), substitua "pub-XXXXXXXXXXXXXXXX" pelo número do seu Publisher ID',
    'Faça deploy e aguarde a verificação do Google (pode levar de 24h a 2 semanas)',
]
for i, s in enumerate(steps_ads, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('6.3. Posições de Anúncios Configuradas', level=2)
doc.add_paragraph(
    'O site possui 5 posições de anúncio otimizadas para máxima receita sem prejudicar a experiência:'
)
ads = [
    ('Ad Top (Leaderboard)', 'Banner horizontal no topo — formato 728×90 desktop / responsivo mobile. data-ad-format="horizontal"'),
    ('Ad Left (Skyscraper)', 'Banner vertical lateral esquerda — formato 160×600. Visível apenas em desktop (>1200px). data-ad-format="vertical"'),
    ('Ad Right (Skyscraper)', 'Banner vertical lateral direita — formato 160×600. Visível apenas em desktop (>1200px). data-ad-format="vertical"'),
    ('Ad In-Article (Fluid)', 'Anúncio nativo entre o resultado do cálculo e o FAQ — zona de alta atenção. data-ad-format="fluid", data-ad-layout="in-article"'),
    ('Ad Bottom (Leaderboard)', 'Banner horizontal no rodapé — formato 728×90 desktop / responsivo mobile. data-ad-format="horizontal"'),
]
t_ads = doc.add_table(rows=len(ads) + 1, cols=2)
t_ads.style = 'Light Grid Accent 1'
t_ads.alignment = WD_TABLE_ALIGNMENT.CENTER
t_ads.rows[0].cells[0].text = 'Posição'
t_ads.rows[0].cells[1].text = 'Descrição e Formato'
for i, (pos, desc) in enumerate(ads, 1):
    t_ads.rows[i].cells[0].text = pos
    t_ads.rows[i].cells[1].text = desc

doc.add_heading('6.4. Arquivo ads.txt', level=2)
doc.add_paragraph(
    'O arquivo ads.txt na raiz do site é obrigatório para que o AdSense verifique a propriedade. '
    'Ele deve ser acessível em: https://ipva.fsncompany.com.br/ads.txt. '
    'Conteúdo: "google.com, pub-XXXXXXXXXXXXXXXX, DIRECT, f08c47fec0942fa0" '
    '(substituir pelo seu Publisher ID).'
)

doc.add_heading('6.5. Consentimento de Cookies (LGPD)', level=2)
doc.add_paragraph(
    'O site inclui um banner de consentimento de cookies que aparece na primeira visita. '
    'É obrigatório pela LGPD (Lei Geral de Proteção de Dados) brasileira e pelas políticas do AdSense.'
)
lgpd_items = [
    'Se o usuário aceitar: anúncios personalizados são exibidos (maior receita)',
    'Se o usuário recusar: anúncios genéricos são exibidos (adsbygoogle.requestNonPersonalizedAds = 1)',
    'O consentimento é salvo em localStorage e não é solicitado novamente',
    'Evento "cookie_consent" é enviado ao GA4 para métricas de taxa de aceite',
]
for item in lgpd_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.6. Placeholders de IDs para Substituir', level=2)
doc.add_paragraph(
    'Ao receber sua aprovação do AdSense, substitua os seguintes placeholders nos arquivos:'
)
placeholders = [
    ('ca-pub-XXXXXXXXXXXXXXXX', 'index.html (6 ocorrências)', 'Seu Publisher ID do AdSense'),
    ('TOP_AD_SLOT_ID', 'index.html', 'Slot ID do anúncio do topo'),
    ('LEFT_AD_SLOT_ID', 'index.html', 'Slot ID do anúncio esquerdo'),
    ('RIGHT_AD_SLOT_ID', 'index.html', 'Slot ID do anúncio direito'),
    ('IN_ARTICLE_AD_SLOT_ID', 'index.html', 'Slot ID do anúncio in-article'),
    ('BOTTOM_AD_SLOT_ID', 'index.html', 'Slot ID do anúncio do rodapé'),
    ('pub-XXXXXXXXXXXXXXXX', 'ads.txt', 'Seu Publisher ID (sem "ca-")'),
]
t_ph = doc.add_table(rows=len(placeholders) + 1, cols=3)
t_ph.style = 'Light Grid Accent 1'
t_ph.alignment = WD_TABLE_ALIGNMENT.CENTER
t_ph.rows[0].cells[0].text = 'Placeholder'
t_ph.rows[0].cells[1].text = 'Arquivo'
t_ph.rows[0].cells[2].text = 'O que colocar'
for i, (ph, arq, desc) in enumerate(placeholders, 1):
    t_ph.rows[i].cells[0].text = ph
    t_ph.rows[i].cells[1].text = arq
    t_ph.rows[i].cells[2].text = desc

doc.add_heading('6.7. Estimativa de Receita', level=2)
doc.add_paragraph(
    'A receita depende do tráfego e do nicho. Para um site de calculadora financeira no Brasil, '
    'a estimativa média é:'
)
revenue = [
    '1.000 visitas/dia → R$ 30-80/mês',
    '5.000 visitas/dia → R$ 150-400/mês',
    '10.000 visitas/dia → R$ 300-900/mês',
    '50.000 visitas/dia → R$ 1.500-5.000/mês',
]
for r in revenue:
    doc.add_paragraph(r, style='List Bullet')
doc.add_paragraph(
    'O nicho de IPVA/finanças tem CPM (custo por mil impressões) acima da média, '
    'especialmente durante o período de pagamento do IPVA (janeiro-março).'
)

doc.add_paragraph()

# --- 7. Observabilidade ---
doc.add_heading('7. Observabilidade e Métricas', level=1)

doc.add_heading('7.1. Google Analytics 4', level=2)
doc.add_paragraph(
    'O sistema está pré-configurado com GA4. Para ativar, substitua "G-XXXXXXXXXX" '
    'no index.html pelo seu Measurement ID obtido em analytics.google.com.'
)

doc.add_heading('7.2. Microsoft Clarity', level=2)
doc.add_paragraph(
    'Ferramenta gratuita da Microsoft para heatmaps e gravação de sessões. '
    'Substitua "CLARITY_ID" no index.html pelo ID do projeto criado em clarity.microsoft.com.'
)

doc.add_heading('7.3. Eventos Rastreados', level=2)
events = [
    ('page_view', 'Acesso à página'),
    ('state_selected', 'Usuário selecionou um estado'),
    ('calculation_performed', 'Cálculo realizado (inclui estado, valor, peso, resultados)'),
    ('recalculate_clicked', 'Botão "Refazer Cálculo" clicado'),
    ('share_clicked', 'Botão "Compartilhar" clicado (método nativo ou clipboard)'),
    ('faq_opened', 'Pergunta do FAQ aberta'),
    ('cookie_consent', 'Consentimento de cookies aceito ou recusado'),
]
t3 = doc.add_table(rows=len(events) + 1, cols=2)
t3.style = 'Light Grid Accent 1'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
t3.rows[0].cells[0].text = 'Evento'
t3.rows[0].cells[1].text = 'Descrição'
for i, (ev, desc) in enumerate(events, 1):
    t3.rows[i].cells[0].text = ev
    t3.rows[i].cells[1].text = desc

doc.add_paragraph()

# --- 8. Hospedagem ---
doc.add_heading('8. Hospedagem e Domínio', level=1)

doc.add_heading('8.1. Render (Static Site)', level=2)
doc.add_paragraph(
    'O projeto é hospedado como Static Site no Render. O arquivo render.yaml contém '
    'toda a configuração necessária, incluindo headers de segurança e cache.'
)

doc.add_heading('8.2. Domínio: ipva.fsncompany.com.br', level=2)
doc.add_paragraph(
    '"ipva" é um subdomínio de "fsncompany.com.br". O processo é:'
)
steps = [
    'Registrar o domínio "fsncompany.com.br" em um registrador (ex: Registro.br)',
    'No Render Dashboard → Custom Domains → Adicionar "ipva.fsncompany.com.br"',
    'No painel DNS do registrador, criar registro CNAME: "ipva" → "seu-app.onrender.com"',
    'Verificar no Render — certificado SSL é emitido automaticamente',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('8.3. Alternativas de Domínio', level=2)
alts = [
    'calculadoraipva.com.br',
    'ipva.calculei.com.br',
    'ipva.simulador.com.br',
    'fsncompany.com.br (usando path /ipva ao invés de subdomínio)',
]
for a in alts:
    doc.add_paragraph(a, style='List Bullet')

doc.add_paragraph()

# --- 9. Deploy ---
doc.add_heading('9. Como Fazer Deploy', level=1)
deploy_steps = [
    'Fazer push do código para o repositório GitHub: github.com/nascimentofe/ipva-calculator',
    'No Render, criar novo Static Site conectado ao repositório GitHub',
    'Branch: main | Publish Directory: . (raiz)',
    'O Render fará deploy automático a cada push no main',
    'Configurar domínio customizado conforme seção 8.2',
]
for s in deploy_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph()

# --- 10. Sobre a PEC ---
doc.add_heading('10. Sobre a PEC do IPVA', level=1)
doc.add_paragraph(
    'A PEC (Proposta de Emenda à Constituição) propõe as seguintes mudanças no IPVA:'
)
pec_items = [
    'Teto de 1%: Garante que o IPVA nunca ultrapassará 1% do valor do veículo',
    'Cálculo por Peso: Veículos mais leves pagam proporcionalmente menos, promovendo justiça tributária',
    'Responsabilidade Fiscal: Compensação através do corte de privilégios',
    'A PEC ainda está em tramitação e os valores apresentados são projeções',
]
for item in pec_items:
    doc.add_paragraph(item, style='List Bullet')

# --- Footer ---
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('© 2026 FSN Company — Documento gerado automaticamente')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(148, 163, 184)

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'IPVA_Calculator_Documentacao.docx')
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
